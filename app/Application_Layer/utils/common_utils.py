from docx import Document
import re
from app.Data_Layer.repository.paths import getPaths

base_path, uploads_path, outputs_path = getPaths()
def extract_table_name(sql_script):
    # Regular expression pattern to match CREATE TABLE statement
    pattern = r'CREATE\s+TABLE\s+(?:\[dbo\]\.)?\[?([^\]\s]+)\]?'
    
    # Find the table name using regular expression
    match = re.search(pattern, sql_script, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    else:
        return None


def doc_text(doc_list,table):
    doc = Document()
    for text in doc_list:
    # Split each text block by newline character to handle each part separately
        for line in text.split("\n"):
            if "SP_Name" in line:
                doc.add_heading(line, level=2)
            else:
                doc.add_paragraph(line)
        doc.add_paragraph("\n")
    doc.save(f"{outputs_path}/impact_files/Impact_analysis_{table[0]}_Document.docx")