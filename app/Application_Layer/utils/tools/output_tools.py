import os
from docx import Document
import pandas as pd


def output_tool(data: list, format_type: str, file_name: str) -> str:
    
    from app.Data_Layer.repository.paths import getPaths

    base_path, uploads_path, outputs_path = getPaths()
    
    """
    Generates output files (CSV or DOCX) from the given data.
    
    Params:
        data (list): A list of dictionaries with the data to be output.
        format_type (str): Desired output format ('csv' or 'docx').
        file_name (str): Desired output file name.
    
    Output:
        A string indicating the file path of the generated output.
        Example: "DOCX file created at /path/to/output.docx"
    """
    # if format_type == "csv":
    #     output_file = os.path.join(outputs_path, "impact_analysis.csv")
    #     df = pd.DataFrame(data)
    #     df.to_csv(output_file, index=False)
    #     return f"CSV file created at {output_file}"
    
    # elif format_type == "docx":
    #     output_file = os.path.join(outputs_path, "TDD_documentation.docx")
    #     doc = Document()
        
    #     for entry in data:
    #         for table_name, content in entry.items():
    #             # Add the main table nam
    #             # e as a heading
    #             doc.add_heading(f'Table: {table_name}', level=1)
                
    #             if isinstance(content, dict):
    #                 # For each parameter inside the main content, create a subheading and its own table
    #                 for param, param_values in content.items():
    #                     if len(param_values.items())>=1:
    #                         # Add parameter as a subheading
    #                         doc.add_heading(f'{param}', level=2)
                            
    #                         if isinstance(param_values, dict):
    #                             # Create a table with sub-parameters and their values
    #                             table = doc.add_table(rows=0, cols=2)
    #                             table.autofit = True
                                
    #                             # Populate the table with the sub-parameters and their values
    #                             for sub_param, value in param_values.items():
    #                                 row_cells = table.add_row().cells
    #                                 row_cells[0].text = str(sub_param)
    #                                 row_cells[1].text = str(value)
                            
    #                         else:
    #                             # Handle non-dictionary values, simply display the value in a paragraph
    #                             doc.add_paragraph(f'Value: {param_values}')
                
    #             elif isinstance(content, list):
    #                 # Handle list content (if any)
    #                 for item in content:
    #                     doc.add_paragraph(str(item))
        
    #     doc.save(output_file)
    return f"empty output tool"

def csv_generator(data: list)->str:
    from app.Data_Layer.repository.paths import getPaths
    base_path, uploads_path, outputs_path = getPaths()
    output_file = os.path.join(outputs_path, "impact_analysis.csv")
    df = pd.DataFrame(data)
    df.to_csv(output_file, index=False)
    return f"CSV file created at {output_file}"

def docx_generator(data: list)->str:
    from app.Data_Layer.repository.paths import getPaths
    base_path, uploads_path, outputs_path = getPaths()
    output_file = os.path.join(outputs_path, "Output.docx")
    doc = Document()
    for entry in data:
        for line in entry.split("\n"):
            doc.add_paragraph(line)
    doc.save(output_file)
    return f"DOCX file created at {output_file}"

def TDD_doc_generator(data: list)->str:
    from app.Data_Layer.repository.paths import getPaths
    base_path, uploads_path, outputs_path = getPaths()
    output_file = os.path.join(outputs_path, "TDD_documentation.docx")
    doc = Document()
    
    for entry in data:
        for table_name, content in entry.items():
            # Add the main table nam
            # e as a heading
            doc.add_heading(f'Table: {table_name}', level=1)
            
            if isinstance(content, dict):
                # For each parameter inside the main content, create a subheading and its own table
                for param, param_values in content.items():
                    if len(param_values.items())>=1:
                        # Add parameter as a subheading
                        doc.add_heading(f'{param}', level=2)
                        
                        if isinstance(param_values, dict):
                            # Create a table with sub-parameters and their values
                            table = doc.add_table(rows=0, cols=2)
                            table.autofit = True
                            
                            # Populate the table with the sub-parameters and their values
                            for sub_param, value in param_values.items():
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(sub_param)
                                row_cells[1].text = str(value)
                        
                        else:
                            # Handle non-dictionary values, simply display the value in a paragraph
                            doc.add_paragraph(f'Value: {param_values}')
            
            elif isinstance(content, list):
                # Handle list content (if any)
                for item in content:
                    doc.add_paragraph(str(item))
    
    doc.save(output_file)
    return f"DOCX file created at {output_file}"